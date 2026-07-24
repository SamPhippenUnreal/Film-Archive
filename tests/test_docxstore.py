"""Tests for the TXT-first Writing store and legacy DOCX compatibility.

Exercises the HTML<->docx round trip (text formatting, bullet lists, embedded
images), the folder CRUD (a document IS a .docx on disk), and the backup that
stages a durable local copy, verifies the linked folder holds it, then purges —
including crash recovery via reconcile().
"""
import io
import os
import tempfile
import unittest

from docx import Document

from app import docxstore
from app.docxstore import (DocxStore, html_to_docx_bytes, docx_bytes_to_html)


def _png_bytes():
    from PIL import Image
    im = Image.new("RGB", (8, 8), (200, 120, 60))
    buf = io.BytesIO()
    im.save(buf, format="PNG")
    return buf.getvalue()


class TestHtmlDocxRoundTrip(unittest.TestCase):
    def test_text_and_formatting_round_trip(self):
        html = ("<div>Hello <b>bold</b> and <i>italic</i> and <u>under</u>.</div>"
                "<ul><li>one</li><li>two</li></ul>")
        data = html_to_docx_bytes(html)
        # it is a real, openable .docx
        doc = Document(io.BytesIO(data))
        text = " ".join(p.text for p in doc.paragraphs)
        self.assertIn("Hello", text)
        self.assertIn("bold", text)
        self.assertIn("one", text)
        self.assertIn("two", text)
        # bold/italic/underline survive as runs
        runs = [r for p in doc.paragraphs for r in p.runs]
        self.assertTrue(any(r.bold and "bold" in r.text for r in runs))
        self.assertTrue(any(r.italic and "italic" in r.text for r in runs))
        self.assertTrue(any(r.underline and "under" in r.text for r in runs))
        # and back to html, the list and text return
        back = docx_bytes_to_html(data)
        self.assertIn("<li>", back)
        self.assertIn("bold", back)

    def test_embedded_image_round_trips_to_data_uri(self):
        png = _png_bytes()
        html = '<div>see <img src="/image/abc"></div>'
        data = html_to_docx_bytes(html, {"/image/abc": (png, "png")})
        # the image is embedded in the docx package
        doc = Document(io.BytesIO(data))
        image_parts = [p for p in doc.part.package.iter_parts()
                       if "image" in (p.content_type or "")]
        self.assertTrue(image_parts, "expected an embedded image part")
        # reading it back yields an inline data: URI image
        back = docx_bytes_to_html(data)
        self.assertIn("<img", back)
        self.assertIn("data:image", back)

    def test_data_uri_image_embeds_without_a_resolver(self):
        import base64
        png = _png_bytes()
        uri = "data:image/png;base64," + base64.b64encode(png).decode()
        data = html_to_docx_bytes(f'<div><img src="{uri}"></div>')
        doc = Document(io.BytesIO(data))
        image_parts = [p for p in doc.part.package.iter_parts()
                       if "image" in (p.content_type or "")]
        self.assertTrue(image_parts)


class TestImageDownscaling(unittest.TestCase):
    """A full-resolution photograph embedded verbatim makes a many-megabyte
    .docx that every save rewrites and every open re-encodes; embedding is
    capped to a document-appropriate size instead."""

    def _jpeg(self, w, h):
        from PIL import Image
        im = Image.new("RGB", (w, h), (120, 160, 90))
        buf = io.BytesIO()
        im.save(buf, format="JPEG", quality=95)
        return buf.getvalue()

    def _embedded(self, data):
        from PIL import Image
        doc = Document(io.BytesIO(data))
        for p in doc.part.related_parts.values():
            if "image" in (getattr(p, "content_type", "") or ""):
                return Image.open(io.BytesIO(p.blob)), p.content_type
        return None, None

    def test_oversized_image_is_capped(self):
        import base64
        uri = ("data:image/jpeg;base64,"
               + base64.b64encode(self._jpeg(3200, 2000)).decode())
        data = html_to_docx_bytes(f'<div><img src="{uri}"></div>')
        im, _ = self._embedded(data)
        self.assertLessEqual(max(im.size), docxstore.EMBED_MAX_DIM)
        # aspect ratio is preserved
        self.assertAlmostEqual(im.size[0] / im.size[1], 3200 / 2000, places=2)

    def test_within_cap_image_is_left_untouched(self):
        # a small image is embedded verbatim — not resized, not recompressed to
        # a different format (so no needless generation loss)
        import base64
        uri = ("data:image/png;base64,"
               + base64.b64encode(_png_bytes()).decode())
        data = html_to_docx_bytes(f'<div><img src="{uri}"></div>')
        im, ct = self._embedded(data)
        self.assertEqual(im.size, (8, 8))
        self.assertIn("png", ct)

    def test_downscaling_is_idempotent(self):
        # re-saving an already-capped document does not shrink it further
        import base64
        uri = ("data:image/jpeg;base64,"
               + base64.b64encode(self._jpeg(3200, 2000)).decode())
        once = html_to_docx_bytes(f'<div><img src="{uri}"></div>')
        twice = html_to_docx_bytes(docx_bytes_to_html(once))
        self.assertEqual(self._embedded(once)[0].size,
                         self._embedded(twice)[0].size)


class TestDocxStoreCrud(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.folder = os.path.join(self.tmp, "docs")
        self.backup = os.path.join(self.tmp, "writing_backup")
        self.store = DocxStore(self.folder, local_dir=self.backup)

    def _docx_files(self):
        return sorted(n for n in os.listdir(self.folder)
                      if n.lower().endswith(".docx"))

    def _writing_files(self):
        return sorted(n for n in os.listdir(self.folder)
                      if n.lower().endswith((".txt", ".docx")) and
                      n != "README.txt")

    def test_create_writes_a_plain_text_file(self):
        doc = self.store.create_doc({"content": "<div>first note</div>"})
        self.assertTrue(doc["filename"].endswith(".txt"))
        self.assertIn(doc["filename"], self._writing_files())
        with open(os.path.join(self.folder, doc["filename"]), encoding="utf-8") as f:
            self.assertEqual(f.read(), "first note")
        self.assertIn("first note", doc["content"])

    def test_save_updates_the_same_file_and_delete_removes_it(self):
        doc = self.store.create_doc({"content": "<div>draft</div>"})
        did = doc["id"]
        self.store.save_doc(did, {"content": "<div>revised text</div>"})
        got = self.store.get_doc(did)
        self.assertIn("revised text", got["content"])
        self.assertEqual(len(self._writing_files()), 1)  # same file, not a new one
        self.store.delete_doc(did)
        self.assertEqual(self._writing_files(), [])
        self.assertIsNone(self.store.get_doc(did))

    def test_rename_moves_document_and_companion_state_without_losing_content(self):
        doc = self.store.create_doc({
            "content": "<div>rename me</div>",
            "editor_state": {
                "content": "<div>rename me</div>",
                "rating": 4,
                "tags": "kept",
            },
        })
        old_path = os.path.join(self.folder, doc["filename"])
        old_state = self.store._state_path(doc["filename"])

        renamed, error = self.store.rename_doc(doc["id"], "New Title")

        self.assertIsNone(error)
        self.assertEqual(renamed["id"], "New Title")
        self.assertEqual(renamed["filename"], "New Title.txt")
        self.assertFalse(os.path.exists(old_path))
        self.assertFalse(os.path.exists(old_state))
        self.assertTrue(os.path.exists(os.path.join(self.folder, "New Title.txt")))
        self.assertTrue(os.path.exists(self.store._state_path("New Title.txt")))
        self.assertIn("rename me", renamed["content"])
        self.assertEqual(renamed["rating"], 4)
        self.assertEqual(renamed["tags"], "kept")

    def test_rename_rejects_conflicts_and_unsupported_titles(self):
        first = self.store.create_doc({"title": "First"})
        self.store.create_doc({"title": "Second"})

        conflict, conflict_error = self.store.rename_doc(first["id"], "Second")
        invalid, invalid_error = self.store.rename_doc(first["id"], "bad/name")

        self.assertIsNone(conflict)
        self.assertIn("already exists", conflict_error)
        self.assertIsNone(invalid)
        self.assertIn("unsupported", invalid_error)
        self.assertTrue(os.path.exists(os.path.join(self.folder, "First.txt")))

    def test_dropped_in_docx_appears_as_a_document(self):
        # a file authored elsewhere, dropped straight into the folder
        d = Document()
        d.add_paragraph("written in Word")
        d.save(os.path.join(self.folder, "From Word.docx"))
        docs = self.store.list_docs()
        titles = {x["title"] for x in docs}
        self.assertIn("From Word", titles)
        match = next(x for x in docs if x["title"] == "From Word")
        self.assertTrue(match["preview_pending"])
        self.assertIn("written in Word",
                      self.store.get_doc(match["id"])["content"])

    def test_unique_names_avoid_clobbering(self):
        a = self.store.create_doc({"content": "<div>a</div>"})
        b = self.store.create_doc({"content": "<div>b</div>"})
        self.assertNotEqual(a["filename"], b["filename"])
        self.assertEqual(len(self._writing_files()), 2)

    def test_save_returns_light_metadata_but_file_has_content(self):
        # the editor discards the save response, so save_doc skips the expensive
        # re-read+reparse and returns light metadata — the document on disk still
        # holds the new content, readable back through get_doc
        did = self.store.create_doc({"content": "<div>alpha</div>"})["id"]
        res = self.store.save_doc(did, {"content": "<div>beta text</div>"})
        self.assertEqual(res["content"], "")
        self.assertEqual(res["id"], did)
        self.assertIn("beta text", self.store.get_doc(did)["content"])

    def test_read_reflects_a_file_changed_on_disk(self):
        # the docx->html shape is cached, but keyed by the file's mtime+size, so
        # a document rewritten outside the app (in Word, or by another instance)
        # is re-read rather than served stale from the cache
        path = os.path.join(self.folder, "Legacy.docx")
        docxstore._atomic_write_bytes(path, html_to_docx_bytes("<div>one</div>"))
        did = "Legacy"
        self.assertIn("one", self.store.get_doc(did)["content"])   # caches it
        import time
        time.sleep(0.02)
        docxstore._atomic_write_bytes(
            path,
            html_to_docx_bytes("<div>a wholly different body</div>"))
        self.assertIn("wholly different", self.store.get_doc(did)["content"])

    def test_editor_state_round_trips_image_groups_and_annotations(self):
        state = {
            "content": '<div>before</div><div class="doc-group" data-gid="g1"></div>',
            "groups": {"g1": {"images": ["photo-a", "photo-b"],
                              "layout": 2, "scale": .78, "bw": False}},
            "annotations": {"cellSize": 2, "cells": [[1, 2, "#334455"]],
                            "wiggly": []},
            "rating": 4,
            "tags": "draft, portrait",
        }
        data = {"content": '<div>before</div><div><img src="/image/photo-a">'
                           '<img src="/image/photo-b"></div>',
                "_images": {"/image/photo-a": (_png_bytes(), "png"),
                            "/image/photo-b": (_png_bytes(), "png")},
                "editor_state": state}
        doc = self.store.create_doc(data)
        with open(os.path.join(self.folder, doc["filename"]), "rb") as f:
            portable = f.read()
        self.assertEqual(portable, b"before")
        self.assertNotIn(_png_bytes(), portable)
        reopened = self.store.get_doc(doc["id"])
        self.assertEqual(reopened["content"], state["content"])
        self.assertEqual(reopened["groups"], state["groups"])
        self.assertEqual(reopened["annotations"], state["annotations"])
        self.assertEqual(reopened["rating"], 4)
        self.assertEqual(reopened["tags"], "draft, portrait")

        # Re-saving the same logical placement stays one managed group rather
        # than compounding raw images and page content on every reopen.
        self.store.save_doc(doc["id"], data)
        again = self.store.get_doc(doc["id"])
        self.assertEqual(again["content"].count('data-gid="g1"'), 1)
        self.assertEqual(len(again["groups"]), 1)

    def test_external_docx_change_invalidates_stale_editor_state(self):
        path = os.path.join(self.folder, "Legacy.docx")
        docxstore._atomic_write_bytes(path, html_to_docx_bytes("<div>original</div>"))
        fields = {"content": "<div>managed</div>",
                  "editor_state": {"content": "<div>managed</div>",
                                   "groups": {"g1": {"images": ["p"]}}}}
        self.store.save_doc("Legacy", fields)
        docxstore._atomic_write_bytes(path, html_to_docx_bytes("<div>external</div>"))
        reopened = self.store.get_doc("Legacy")
        self.assertIn("external", reopened["content"])
        self.assertEqual(reopened["groups"], {})

    def test_external_text_change_invalidates_stale_editor_state(self):
        doc = self.store.create_doc({
            "content": "<div>managed</div>",
            "editor_state": {"content": "<div>managed</div>",
                             "groups": {"g1": {"images": ["p"]}}},
        })
        path = os.path.join(self.folder, doc["filename"])
        docxstore._atomic_write_bytes(path, b"external plain text")
        reopened = self.store.get_doc(doc["id"])
        self.assertIn("external plain text", reopened["content"])
        self.assertEqual(reopened["groups"], {})

    def test_archive_list_does_not_parse_external_docx_eagerly(self):
        from unittest import mock
        path = os.path.join(self.folder, "Large legacy.docx")
        docxstore._atomic_write_bytes(path, html_to_docx_bytes("<div>legacy</div>"))
        with mock.patch("app.docxstore.docx_bytes_to_html",
                        side_effect=AssertionError("eager parse")):
            docs = self.store.list_docs()
        self.assertEqual(len(docs), 1)
        self.assertTrue(docs[0]["preview_pending"])

    def test_duplicate_and_delete_follow_editor_state(self):
        doc = self.store.create_doc({
            "content": "<div>body</div>",
            "editor_state": {"content": "<div>body</div>", "rating": 3},
        })
        copy = self.store.duplicate_doc(doc["id"])
        self.assertEqual(copy["rating"], 3)
        self.assertTrue(os.path.exists(self.store._state_path(copy["filename"])))
        self.store.delete_doc(copy["id"])
        self.assertFalse(os.path.exists(self.store._state_path(copy["filename"])))


class TestWritingBackup(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.mkdtemp()
        self.folder = os.path.join(self.tmp, "docs")
        self.backup_dir = os.path.join(self.tmp, "writing_backup")

    def test_backup_purges_after_a_confirmed_commit(self):
        store = DocxStore(self.folder, local_dir=self.backup_dir)
        store.create_doc({"content": "<div>safe</div>"})
        # once the linked folder holds the file, the local copy is gone
        journal = os.path.join(self.backup_dir, "journal")
        docs = os.path.join(self.backup_dir, "docs")
        self.assertEqual([n for n in os.listdir(journal)
                          if n.endswith(".json")], [])
        self.assertEqual([n for n in os.listdir(docs)
                          if n.endswith((".docx", ".txt"))], [])

    def test_reconcile_replays_a_staged_file_the_folder_is_missing(self):
        # simulate a crash: bytes staged locally, never written to the folder
        os.makedirs(os.path.join(self.backup_dir, "docs"), exist_ok=True)
        os.makedirs(os.path.join(self.backup_dir, "journal"), exist_ok=True)
        data = html_to_docx_bytes("<div>recovered</div>")
        name = "Recovered.docx"
        docxstore._atomic_write_bytes(
            os.path.join(self.backup_dir, "docs", name), data)
        docxstore._atomic_write_json(
            os.path.join(self.backup_dir, "journal", name + ".json"),
            {"name": name, "checksum": docxstore._sha(data),
             "ts": 0, "state": "pending"})
        # launching the store reconciles: the folder gets the file, backup purges
        store = DocxStore(self.folder, local_dir=self.backup_dir)
        self.assertTrue(os.path.exists(os.path.join(self.folder, name)))
        self.assertEqual([n for n in os.listdir(
            os.path.join(self.backup_dir, "journal")) if n.endswith(".json")], [])
        # and it is readable as a document
        got = store.get_doc("Recovered")
        self.assertIsNotNone(got)
        self.assertIn("recovered", got["content"])


if __name__ == "__main__":
    unittest.main()
