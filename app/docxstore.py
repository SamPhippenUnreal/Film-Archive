"""Writing Mode, folder-backed by plain text with legacy DOCX support.

New Archive documents are ordinary UTF-8 ``.txt`` files.  Exact Archive-only
layout (formatting, placed photographs and annotations) lives in a companion
``.archive-writing`` sidecar, while the text itself stays readable everywhere.
Existing ``.docx`` files remain discoverable and editable; their current DOCX
round-trip is retained for compatibility, but the fast path never rebuilds or
embeds images into a text document.

Safety — the ``writing_backup`` journal
---------------------------------------
The linked folder plays the same role the shared ``cache`` folder plays for the
image archive: it is the authoritative, possibly cloud-synced destination. To
protect against a slow/offline/locked destination or a crash mid-write, every
commit is written **backup-first**, mirroring the image archive's write-ahead
journal (see backup.py):

    writing_backup/
        docs/     <name>.docx        a durable local copy of the bytes
        journal/  <name>.docx.json   {name, checksum, ts, state}

Flow per commit: stage the bytes locally, write them to the linked folder,
verify the folder now holds exactly those bytes, then purge the local copy. A
crash leaves the staged copy in place; on the next launch ``reconcile`` replays
anything the linked folder is missing or that doesn't match, and purges what is
already up to date. The local backup therefore only ever holds work not yet
safely in the linked folder.
"""
import base64
import hashlib
import html as html_lib
import io
import os
import re
import threading
import time
from html.parser import HTMLParser

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from .safeio import (
    atomic_write_bytes as _shared_atomic_write_bytes,
    atomic_write_json as _shared_atomic_write_json,
    is_within,
    read_json as _shared_read_json,
    sha256_bytes,
)
from .writing_state import make_sidecar, read_matching_sidecar, stat_signature

try:
    from PIL import Image
except Exception:                       # pragma: no cover - Pillow is a dep
    Image = None


# ————————————————————————— low-level io helpers —————————————————————————

def _sha(blob):
    return sha256_bytes(blob)


def _atomic_write_bytes(path, data):
    _shared_atomic_write_bytes(path, data)


def _atomic_write_json(path, obj):
    _shared_atomic_write_json(path, obj)


def _read_json(path):
    return _shared_read_json(path, dictionaries_only=False)


class _PlainTextParser(HTMLParser):
    """Small HTML-to-text projection for the portable ``.txt`` source."""

    BREAKS = {"div", "p", "li", "blockquote", "h1", "h2", "h3", "h4",
              "h5", "h6", "br"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts = []

    def _newline(self):
        if self.parts and not self.parts[-1].endswith("\n"):
            self.parts.append("\n")

    def handle_starttag(self, tag, attrs):
        if tag.lower() in self.BREAKS:
            self._newline()
        if tag.lower() == "li":
            self.parts.append("- ")

    def handle_endtag(self, tag):
        if tag.lower() in self.BREAKS:
            self._newline()

    def handle_data(self, data):
        self.parts.append(data)


def html_to_plain_text(value):
    parser = _PlainTextParser()
    try:
        parser.feed(value or "")
        parser.close()
    except Exception:
        return re.sub(r"<[^>]+>", "", value or "").strip()
    text = "".join(parser.parts).replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip("\n")


def plain_text_to_html(value):
    lines = (value or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    return "".join("<div>%s</div>" % (html_lib.escape(line) if line else "<br>")
                   for line in lines)


_UNSAFE = re.compile(r'[<>:"/\\|?*\x00-\x1f]')


def _safe_filename(name):
    """A filesystem-safe document title, without extension."""
    name = _UNSAFE.sub(" ", (name or "").strip()).strip(". ")
    name = re.sub(r"\s+", " ", name)
    return name[:80] or "Untitled"


# ————————————————————————— html  ->  docx —————————————————————————

_BLOCK = {"div", "p", "li", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote"}
_HEADINGS = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}


def _image_width(blob):
    """A page-margin-safe width for an embedded image."""
    if Image is None:
        return Inches(5.0)
    try:
        im = Image.open(io.BytesIO(blob))
        dpi = (im.info.get("dpi") or (96, 96))[0] or 96
        w_in = im.size[0] / float(dpi)
        return Inches(max(1.0, min(6.0, w_in)))
    except Exception:
        return Inches(5.0)


# The longest edge, in pixels, of a photograph embedded into a document. Linked
# photographs are full-resolution originals (commonly 3000–6000px wide); a
# document only ever displays one at page width, so embedding the original
# makes the .docx many megabytes. Every autosave then rebuilds and rewrites that
# whole payload and every open re-encodes it to base64 — the single biggest
# cause of the Writing context feeling slow. Capping the longest edge shrinks a
# typical photo roughly ten-fold with no visible loss at document scale.
EMBED_MAX_DIM = 1600


def _downscale_image(blob):
    """Shrink an oversized embedded photograph to a document-appropriate size.

    Idempotent: an image already within the cap is returned untouched, so
    repeated saves never recompress it (no generation loss) and only the first
    save of an oversized image pays the cost. Any failure returns the original
    bytes, so a quirky image is embedded verbatim rather than dropped."""
    if Image is None:
        return blob
    try:
        im = Image.open(io.BytesIO(blob))
        w, h = im.size
        if max(w, h) <= EMBED_MAX_DIM:
            return blob
        scale = EMBED_MAX_DIM / float(max(w, h))
        im = im.convert("RGB").resize(
            (max(1, round(w * scale)), max(1, round(h * scale))), Image.LANCZOS)
        out = io.BytesIO()
        im.save(out, format="JPEG", quality=85, optimize=True)
        return out.getvalue()
    except Exception:
        return blob


class _HtmlToDocx(HTMLParser):
    """Convert the editor's content HTML into Word paragraphs and runs.

    Supports the formatting the editor produces: block paragraphs, bold /
    italic / underline runs, bullet lists, and inline images (data: URIs, or
    ``/image/<id>`` sources resolved to bytes by the caller)."""

    def __init__(self, document, images):
        super().__init__(convert_charrefs=True)
        self.doc = document
        self.images = images or {}
        self.para = None
        self.bold = self.ital = self.und = 0
        self.list_depth = 0
        self.heading = 0

    def _flush(self):
        self.para = None

    def _ensure(self):
        if self.para is None:
            self.para = self.doc.add_paragraph()
            if self.heading:
                try:
                    self.para.style = self.doc.styles[f"Heading {self.heading}"]
                except KeyError:
                    pass
            elif self.list_depth > 0:
                try:
                    self.para.style = self.doc.styles["List Bullet"]
                except KeyError:
                    pass
        return self.para

    def handle_starttag(self, tag, attrs):
        if tag in ("b", "strong"):
            self.bold += 1
        elif tag in ("i", "em"):
            self.ital += 1
        elif tag == "u":
            self.und += 1
        elif tag in ("ul", "ol"):
            self.list_depth += 1
            self._flush()
        elif tag == "br":
            if self.para is not None:
                self.para.add_run().add_break()
        elif tag == "img":
            self._add_image(dict(attrs).get("src", ""))
        elif tag in _HEADINGS:
            self._flush()
            self.heading = _HEADINGS[tag]
        elif tag in _BLOCK:
            self._flush()

    def handle_startendtag(self, tag, attrs):
        self.handle_starttag(tag, attrs)

    def handle_endtag(self, tag):
        if tag in ("b", "strong"):
            self.bold = max(0, self.bold - 1)
        elif tag in ("i", "em"):
            self.ital = max(0, self.ital - 1)
        elif tag == "u":
            self.und = max(0, self.und - 1)
        elif tag in ("ul", "ol"):
            self.list_depth = max(0, self.list_depth - 1)
            self._flush()
        elif tag in _HEADINGS:
            self.heading = 0
            self._flush()
        elif tag in _BLOCK:
            self._flush()

    def handle_data(self, data):
        if not data or (data.strip() == "" and self.para is None):
            return
        run = self._ensure().add_run(data)
        run.bold = self.bold > 0
        run.italic = self.ital > 0
        run.underline = self.und > 0

    def _add_image(self, src):
        blob = None
        if src.startswith("data:"):
            try:
                blob = base64.b64decode(src.split(",", 1)[1])
            except Exception:
                blob = None
        else:
            entry = self.images.get(src)
            if entry:
                blob = entry[0]
        if not blob:
            return
        blob = _downscale_image(blob)
        try:
            self._ensure().add_run().add_picture(
                io.BytesIO(blob), width=_image_width(blob))
        except Exception:
            pass


def html_to_docx_bytes(html, images=None):
    """Serialize editor content HTML to ``.docx`` bytes."""
    doc = Document()
    parser = _HtmlToDocx(doc, images or {})
    parser.feed(html or "")
    parser.close()
    if not doc.paragraphs:
        doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ————————————————————————— docx  ->  html —————————————————————————

def _esc(text):
    return (text.replace("&", "&amp;").replace("<", "&lt;")
                .replace(">", "&gt;"))


def _run_images(run, part):
    out = []
    for blip in run._element.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if not rid:
            continue
        try:
            image_part = part.related_parts[rid]
            b64 = base64.b64encode(image_part.blob).decode("ascii")
            ct = image_part.content_type or "image/png"
            out.append(f'<img src="data:{ct};base64,{b64}">')
        except Exception:
            pass
    return "".join(out)


def _run_html(run, part):
    imgs = _run_images(run, part)
    text = _esc(run.text or "")
    if text:
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        if run.underline:
            text = f"<u>{text}</u>"
    return imgs + text


def docx_bytes_to_html(data):
    """Render a ``.docx`` file into content HTML the editor can display."""
    doc = Document(io.BytesIO(data))
    part = doc.part
    out = []
    in_list = False
    for para in doc.paragraphs:
        style = (para.style.name or "") if para.style is not None else ""
        is_li = style.startswith("List")
        inner = "".join(_run_html(r, part) for r in para.runs)
        if is_li:
            if not in_list:
                out.append("<ul>")
                in_list = True
            out.append(f"<li>{inner or '<br>'}</li>")
            continue
        if in_list:
            out.append("</ul>")
            in_list = False
        out.append(f"<div>{inner or '<br>'}</div>")
    if in_list:
        out.append("</ul>")
    return "".join(out)


def docx_text(data):
    """Plain text of a document, for previews and search."""
    try:
        doc = Document(io.BytesIO(data))
    except Exception:
        return ""
    return " ".join(p.text for p in doc.paragraphs).strip()


def docx_plain_text(data):
    """Lowercase plain text of a document, for search."""
    return docx_text(data).lower()


# ————————————————————————— the write-ahead backup —————————————————————————

class _DocxBackup:
    """A durable local copy of each committed ``.docx`` until the linked folder
    confirms it, then purged — the binary counterpart of backup.py's journal."""

    def __init__(self, local_dir):
        self.dir = local_dir
        self.docs_dir = os.path.join(local_dir, "docs")
        self.journal_dir = os.path.join(local_dir, "journal")
        self._lock = threading.RLock()
        self.ok = True
        try:
            os.makedirs(self.docs_dir, exist_ok=True)
            os.makedirs(self.journal_dir, exist_ok=True)
        except OSError:
            self.ok = False

    def _jpath(self, name):
        return os.path.join(self.journal_dir, name + ".json")

    def _dpath(self, name):
        return os.path.join(self.docs_dir, name)

    def stage(self, name, data):
        """Durably record the bytes locally, before the linked folder is
        touched."""
        if not self.ok:
            return
        with self._lock:
            try:
                _atomic_write_bytes(self._dpath(name), data)
                _atomic_write_json(self._jpath(name), {
                    "name": name, "checksum": _sha(data),
                    "ts": time.time(), "state": "pending"})
            except OSError:
                pass

    def confirm(self, name, target_dir):
        """The linked folder now holds this file — if it matches the staged
        bytes, the local safety copy has done its job and is purged."""
        if not self.ok:
            return
        with self._lock:
            entry = _read_json(self._jpath(name)) or {}
            try:
                with open(os.path.join(target_dir, name), "rb") as f:
                    cur = f.read()
            except OSError:
                return
            if _sha(cur) == entry.get("checksum"):
                self._purge(name)

    def forget(self, name):
        """Drop any staged copy for a deleted document, so reconcile never
        resurrects it."""
        if not self.ok:
            return
        with self._lock:
            self._purge(name)

    def _purge(self, name):
        for p in (self._dpath(name), self._jpath(name)):
            try:
                os.remove(p)
            except OSError:
                pass

    def reconcile(self, target_dir):
        """On launch: replay anything the linked folder is missing or that does
        not match the staged bytes, then purge what is already up to date."""
        if not self.ok:
            return
        with self._lock:
            journals = []
            for dirpath, _, names in os.walk(self.journal_dir):
                journals.extend(os.path.join(dirpath, name) for name in names
                                if name.endswith(".json") and ".tmp." not in name)
            for journal_path in journals:
                entry = _read_json(journal_path)
                if not isinstance(entry, dict) or not entry.get("name"):
                    continue
                name = entry["name"]
                staged = self._dpath(name)
                if not os.path.exists(staged):
                    self._purge(name)
                    continue
                try:
                    with open(staged, "rb") as f:
                        staged_bytes = f.read()
                except OSError:
                    continue
                if _sha(staged_bytes) != entry.get("checksum"):
                    self._purge(name)          # corrupt/incomplete stage
                    continue
                target = os.path.join(target_dir, name)
                try:
                    with open(target, "rb") as f:
                        cur = f.read()
                except OSError:
                    cur = None
                if cur is not None and _sha(cur) == entry.get("checksum"):
                    self._purge(name)          # already safely landed
                    continue
                try:                            # replay, verify, purge
                    _atomic_write_bytes(target, staged_bytes)
                    with open(target, "rb") as f:
                        if _sha(f.read()) == entry.get("checksum"):
                            self._purge(name)
                except OSError:
                    pass

    def pending(self):
        """How many committed docs are not yet confirmed in the folder."""
        if not self.ok:
            return 0
        try:
            return sum(
                1 for _, _, names in os.walk(self.journal_dir)
                for name in names
                if name.endswith(".json") and ".tmp." not in name)
        except OSError:
            return 0


# ————————————————————————— the store —————————————————————————

_README = (
    "Archive — Writing documents\n"
    "===========================\n\n"
    "New Archive documents are UTF-8 .txt files, readable in any text editor.\n"
    "Formatting, placed photographs and annotations live in the companion\n"
    ".archive-writing folder and appear inside Archive. Existing .docx files\n"
    "remain available for compatibility. A local write-ahead backup keeps\n"
    "each save safe until this folder confirms it.\n"
)


class DocxStore:
    """Mixed ``.txt``/legacy ``.docx`` Writing store."""

    def __init__(self, root_dir, local_dir=None):
        self.dir = root_dir
        os.makedirs(root_dir, exist_ok=True)
        # A .docx is the portable, user-readable document.  A small companion
        # state file preserves Archive-only editing semantics that Word cannot
        # represent (image-group ids/layout, canvas annotations, rating/tags).
        # It lives with the documents so cloud-synced Writing folders retain
        # the exact editor model across machines.
        self.state_dir = os.path.join(root_dir, ".archive-writing")
        os.makedirs(self.state_dir, exist_ok=True)
        self._lock = threading.RLock()
        self._ids = {}                       # id (stem) -> actual filename
        # filename -> ((mtime, size), shaped-dict): the docx->html conversion is
        # the expensive step of every list/open, so its result is memoised and
        # only recomputed when the file on disk actually changes (a fresh save,
        # or a file edited outside the app). Keeps opening the archive cheap.
        self._shape_cache = {}
        self.backup = _DocxBackup(local_dir) if local_dir else None
        if self.backup and self.backup.ok:
            self.backup.reconcile(self.dir)  # crash-safe replay on launch
        self._write_readme()

    # ---- helpers ----------------------------------------------------------

    def _write_readme(self):
        path = os.path.join(self.dir, "README.txt")
        if not os.path.exists(path):
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(_README)
            except OSError:
                pass

    def _scan(self):
        ids = {}
        names = []
        for dirpath, dirnames, filenames in os.walk(self.dir, followlinks=False):
            dirnames[:] = [
                name for name in dirnames
                if name.casefold() not in {
                    ".archive-writing", "cache", "trash", "project canvas"}
                and not os.path.islink(os.path.join(dirpath, name))
            ]
            for filename in filenames:
                relative = os.path.relpath(
                    os.path.join(dirpath, filename), self.dir).replace("\\", "/")
                names.append(relative)
        names.sort(key=lambda value: (
            value.count("/"), value.casefold()))
        for name in names:
            lower = name.casefold()
            basename = os.path.basename(name)
            if (not lower.endswith((".txt", ".docx")) or
                    basename.startswith("~$") or
                    ("/" not in name and basename == "README.txt")):
                continue
            stem = os.path.splitext(basename)[0]
            parent = os.path.dirname(name).replace("\\", "/")
            doc_id = (stem if not parent else
                      f"{stem}~{hashlib.sha1(name.encode('utf-8')).hexdigest()[:8]}")
            if doc_id in ids:
                suffix = "txt" if lower.endswith(".txt") else "docx"
                doc_id = f"{stem}~{suffix}"
            while doc_id in ids:
                doc_id += "2"
            ids[doc_id] = name
        self._ids = ids
        return ids

    def _unique_name(self, title, extension=".txt", parent=""):
        base = _safe_filename(title) if title else "Untitled"
        filename = base + extension
        name = f"{parent}/{filename}" if parent else filename
        i = 2
        existing = {n.lower() for n in self._ids.values()}
        while name.lower() in existing:
            filename = f"{base} {i}{extension}"
            name = f"{parent}/{filename}" if parent else filename
            i += 1
        return name

    def _commit(self, name, data):
        """Backup-first: stage locally, write the linked folder, verify+purge."""
        path = os.path.join(self.dir, name)
        if self.backup:
            self.backup.stage(name, data)
        _atomic_write_bytes(path, data)
        if self.backup:
            self.backup.confirm(name, self.dir)
        self._shape_cache.pop(name, None)    # rewritten: its cached shape is stale

    def _state_path(self, name):
        return os.path.join(self.state_dir, name + ".json")

    def _write_editor_state(self, name, data, state):
        if not isinstance(state, dict):
            return
        clean = make_sidecar(
            data, os.path.splitext(name)[1], state,
            stat_signature(os.path.join(self.dir, name)))
        _atomic_write_json(self._state_path(name), clean)

    def _read_editor_state(self, name, data):
        return read_matching_sidecar(self._state_path(name), data)

    def _state_for_summary(self, name, st):
        state = _read_json(self._state_path(name))
        if not isinstance(state, dict):
            return None
        signature = state.get("source_stat")
        if signature != [st.st_mtime_ns, st.st_size]:
            return None
        return state

    def _light(self, doc_id, name, mt=None):
        """The document's metadata without the costly docx->html conversion —
        enough for the save response (which the editor discards) and to fill in
        a shape whose content is supplied by the caller."""
        if mt is None:
            try:
                mt = os.path.getmtime(os.path.join(self.dir, name))
            except OSError:
                mt = time.time()
        return {
            "id": doc_id, "title": os.path.splitext(os.path.basename(name))[0],
            "filename": os.path.basename(name), "rel_path": name,
            "format": os.path.splitext(name)[1].lower().lstrip("."),
            "content": "",
            # kept for shape-compatibility with the editor's document model
            "groups": {}, "annotations": {}, "rating": 0, "tags": "",
            "created": mt, "updated": mt,
        }

    def _shape(self, doc_id, name):
        path = os.path.join(self.dir, name)
        try:
            st = os.stat(path)
            try:
                sst = os.stat(self._state_path(name))
                state_key = (sst.st_mtime, sst.st_size)
            except OSError:
                state_key = None
            key = (st.st_mtime, st.st_size, state_key)
        except OSError:
            key = None
        cached = self._shape_cache.get(name)
        if cached is not None and key is not None and cached[0] == key:
            return dict(cached[1])           # a copy — callers may reassign keys
        shaped = self._light(doc_id, name, mt=key[0] if key else None)
        try:
            with open(path, "rb") as f:
                data = f.read()
            state = self._read_editor_state(name, data)
            if state is not None:
                shaped["content"] = state.get("content") or ""
                shaped["groups"] = state.get("groups") or {}
                shaped["annotations"] = state.get("annotations") or {}
                shaped["rating"] = state.get("rating") or 0
                shaped["tags"] = state.get("tags") or ""
            elif name.lower().endswith(".txt"):
                shaped["content"] = plain_text_to_html(data.decode("utf-8-sig", "replace"))
            else:
                # A document created or changed outside Archive has no matching
                # companion model.  Parse it normally and bound its inline
                # images in the editor; never trust stale state from older bytes.
                shaped["content"] = docx_bytes_to_html(data)
        except Exception:
            shaped["content"] = ""
        if key is not None:
            self._shape_cache[name] = (key, dict(shaped))
        return shaped

    def _summary(self, doc_id, name):
        path = os.path.join(self.dir, name)
        try:
            st = os.stat(path)
        except OSError:
            return self._light(doc_id, name)
        shaped = self._light(doc_id, name, st.st_mtime)
        state = self._state_for_summary(name, st)
        if state is not None:
            shaped.update({
                "content": state.get("content") or "",
                "groups": state.get("groups") or {},
                "annotations": state.get("annotations") or {},
                "rating": state.get("rating") or 0,
                "tags": state.get("tags") or "",
            })
        elif name.lower().endswith(".txt"):
            try:
                with open(path, "r", encoding="utf-8-sig", errors="replace") as f:
                    shaped["content"] = plain_text_to_html(f.read())
            except OSError:
                pass
        else:
            # Legacy external Word files are parsed only when opened or when a
            # visible preview explicitly asks for them, never on archive load.
            shaped["preview_pending"] = True
        return shaped

    # ---- crud -------------------------------------------------------------

    def list_docs(self):
        if self.backup:
            self.backup.reconcile(self.dir)
        with self._lock:
            ids = dict(self._scan())
        docs = [self._summary(i, n) for i, n in ids.items()]
        docs.sort(key=lambda d: d.get("updated") or 0, reverse=True)
        return docs

    def get_doc(self, doc_id):
        with self._lock:
            self._scan()
            name = self._ids.get(doc_id)
        if not name:
            return None
        return self._shape(doc_id, name)

    def source_path(self, doc_id):
        """Canonical path for a known document id, or ``None``.

        Importers use this rather than joining a request value to the linked
        Writing folder.  Only an id discovered by ``_scan`` can resolve.
        """
        with self._lock:
            self._scan()
            name = self._ids.get(str(doc_id))
        if not name:
            return None
        path = os.path.realpath(os.path.join(self.dir, *name.split("/")))
        if not is_within(self.dir, path, allow_root=False):
            return None
        return path if os.path.isfile(path) else None

    def export_state(self, doc_id):
        """Portable Archive preview state for a Project import."""
        path = self.source_path(doc_id)
        if not path:
            return None
        try:
            with open(path, "rb") as f:
                data = f.read()
        except OSError:
            return None
        shaped = self.get_doc(str(doc_id))
        if not shaped:
            return None
        return {
            "version": 2,
            "source_checksum": _sha(data),
            "source_format": os.path.splitext(path)[1].lower().lstrip("."),
            "content": shaped.get("content") or "",
            "groups": shaped.get("groups") or {},
            "annotations": shaped.get("annotations") or {},
            "rating": shaped.get("rating") or 0,
            "tags": shaped.get("tags") or "",
        }

    def create_doc(self, fields=None):
        fields = fields or {}
        with self._lock:
            self._scan()
            name = self._unique_name(fields.get("title"), ".txt")
            data = html_to_plain_text(fields.get("content", "")).encode("utf-8")
            self._commit(name, data)
            self._write_editor_state(name, data, fields.get("editor_state"))
            self._scan()
            new_id = next((i for i, n in self._ids.items() if n == name),
                          os.path.splitext(name)[0])
        return self._shape(new_id, name)

    def save_doc(self, doc_id, fields):
        fields = fields or {}
        with self._lock:
            self._scan()
            name = self._ids.get(doc_id) or (_safe_filename(doc_id) + ".txt")
            if name.lower().endswith(".txt"):
                data = html_to_plain_text(fields.get("content", "")).encode("utf-8")
            else:
                data = html_to_docx_bytes(fields.get("content", ""),
                                          fields.get("_images"))
            self._commit(name, data)
            self._write_editor_state(name, data, fields.get("editor_state"))
            self._scan()
            # The editor discards the save response, so skip re-reading and
            # re-converting the file we just wrote (a full docx->html of a
            # possibly image-heavy document) — return light metadata only. The
            # next list/open re-parses this one document once and re-caches it.
            return self._light(doc_id, name)

    def rename_doc(self, doc_id, title):
        """Rename a known TXT/DOCX and carry its Archive editor state with it."""
        if not isinstance(title, str) or not title.strip():
            return None, "document title cannot be empty"
        requested = title.strip()
        clean = _safe_filename(requested)
        if clean != requested:
            return None, "that document title contains unsupported characters"
        if clean.split(".", 1)[0].upper() in {
                "CON", "PRN", "AUX", "NUL",
                *(f"COM{i}" for i in range(1, 10)),
                *(f"LPT{i}" for i in range(1, 10))}:
            return None, "that document title is reserved"

        with self._lock:
            self._scan()
            old_name = self._ids.get(str(doc_id))
            if not old_name:
                return None, "document could not be found"
            extension = os.path.splitext(old_name)[1].lower()
            parent = os.path.dirname(old_name).replace("\\", "/")
            new_name = f"{parent}/{clean}{extension}" if parent else clean + extension
            if new_name == old_name:
                return self._shape(str(doc_id), old_name), None
            if any(name.casefold() == new_name.casefold() and name != old_name
                   for name in self._ids.values()):
                return None, "a document with that title already exists"

            old_path = os.path.join(self.dir, old_name)
            new_path = os.path.join(self.dir, new_name)
            old_state = self._state_path(old_name)
            new_state = self._state_path(new_name)
            if (os.path.normcase(new_path) != os.path.normcase(old_path) and
                    os.path.exists(new_path)):
                return None, "a document with that title already exists"

            state_moved = False
            try:
                # A temporary hop makes case-only renames deterministic on
                # case-insensitive filesystems.
                if os.path.normcase(new_path) == os.path.normcase(old_path):
                    temporary = old_path + f".rename.{os.getpid()}"
                    os.replace(old_path, temporary)
                    os.replace(temporary, new_path)
                else:
                    os.rename(old_path, new_path)
                if os.path.exists(old_state):
                    os.replace(old_state, new_state)
                    state_moved = True
            except OSError:
                if state_moved and os.path.exists(new_state):
                    try:
                        os.replace(new_state, old_state)
                    except OSError:
                        pass
                if os.path.exists(new_path) and not os.path.exists(old_path):
                    try:
                        os.replace(new_path, old_path)
                    except OSError:
                        pass
                return None, "document could not be renamed"

            if self.backup:
                self.backup.forget(old_name)
            self._shape_cache.pop(old_name, None)
            self._shape_cache.pop(new_name, None)
            self._scan()
            new_id = next((key for key, name in self._ids.items()
                           if name == new_name), clean)
            return self._shape(new_id, new_name), None

    def duplicate_doc(self, doc_id):
        with self._lock:
            self._scan()
            name = self._ids.get(doc_id)
            if not name:
                return None
            try:
                with open(os.path.join(self.dir, name), "rb") as f:
                    data = f.read()
            except OSError:
                return None
            extension = os.path.splitext(name)[1].lower()
            parent = os.path.dirname(name).replace("\\", "/")
            title = os.path.splitext(os.path.basename(name))[0]
            new_name = self._unique_name(title, extension, parent)
            self._commit(new_name, data)
            state = self._read_editor_state(name, data)
            if state is not None:
                self._write_editor_state(new_name, data, state)
            self._scan()
            new_id = next((i for i, n in self._ids.items() if n == new_name),
                          os.path.splitext(new_name)[0])
        return self._shape(new_id, new_name)

    def delete_doc(self, doc_id):
        with self._lock:
            self._scan()
            name = self._ids.get(doc_id)
            if name:
                try:
                    os.remove(os.path.join(self.dir, name))
                except OSError:
                    pass
                if self.backup:
                    self.backup.forget(name)
                self._ids.pop(doc_id, None)
                self._shape_cache.pop(name, None)
                try:
                    os.remove(self._state_path(name))
                except OSError:
                    pass
        return True
